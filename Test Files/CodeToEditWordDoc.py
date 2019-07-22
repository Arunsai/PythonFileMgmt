from docx import Document

document=Document("My_WordDoc_1.docx")

document.add_heading('The Python code created Document')

paragraph = document.add_paragraph("This is a test and first paragraph created for this project. To ensure that there is enough content in this paragraph am adding the famous sample text next to it. Lorem ipsum dolor sit amet, consectetur adipiscing elit, sed do eiusmod tempor incididunt ut labore et dolore magna aliqua. Ut enim ad minim veniam, quis nostrud exercitation ullamco laboris nisi ut aliquip ex ea commodo consequat. Duis aute irure dolor in reprehenderit in voluptate velit esse cillum dolore eu fugiat nulla pariatur. Excepteur sint occaecat cupidatat non proident, sunt in culpa qui officia deserunt mollit anim id est laborum.")

document.add_heading('This is a level 2 heading',level=2)

paragraph2 = document.add_paragraph('This is paragraph 2 and followed by sample test. Sit amet justo donec enim diam vulputate ut pharetra. Commodo quis imperdiet massa tincidunt nunc pulvinar sapien et. Elementum nibh tellus molestie nunc non blandit massa. Neque laoreet suspendisse interdum consectetur libero id faucibus nisl. Enim tortor at auctor urna nunc id cursus. Neque convallis a cras semper auctor neque vitae tempus quam. Congue mauris rhoncus aenean vel. Est lorem ipsum dolor sit amet consectetur adipiscing. Mauris pharetra et ultrices neque ornare aenean. Proin fermentum leo vel orci porta.')

document.add_page_break()

document.add_heading('This is also a level 2 heading',level=2)

paragraph3 = document.add_paragraph('This is paragraph 3 which should be followed by a line break.. So assuming this appears in a new page. The  remaining part of the paragraph is from the famous Lorum Ipsum. Mattis aliquam faucibus purus in massa. Mi bibendum neque egestas congue quisque egestas diam in. Donec adipiscing tristique risus nec feugiat in. Nunc sed id semper risus. Sed viverra tellus in hac. Consequat ac felis donec et odio pellentesque diam volutpat. Non sodales neque sodales ut etiam. Vitae elementum curabitur vitae nunc sed velit. Quisque id diam vel quam elementum pulvinar etiam. Integer vitae justo eget magna fermentum iaculis eu. Risus quis varius quam quisque id diam. Egestas maecenas pharetra convallis posuere morbi leo urna molestie.')

document.save('CodeOutputs\My_WordDoc_2.docx')